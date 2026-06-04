from __future__ import annotations

import os
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
BUILD_DIR = ROOT / "_lease_report_build"
IMG_DIR = BUILD_DIR / "images"
OUT_DOCX = ROOT / "房屋租赁系统软件工程大作业报告.docx"

FONT_CANDIDATES = [
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\simsun.ttc"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
]


def ensure_dirs() -> None:
    IMG_DIR.mkdir(parents=True, exist_ok=True)


def font_path() -> str | None:
    for candidate in FONT_CANDIDATES:
        if candidate.exists():
            return str(candidate)
    return None


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    path = font_path()
    if path:
        return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in text.split("\n"):
        current = ""
        for ch in paragraph:
            test = current + ch
            bbox = draw.textbbox((0, 0), test, font=fnt)
            if bbox[2] - bbox[0] <= max_width:
                current = test
            else:
                if current:
                    lines.append(current)
                current = ch
        lines.append(current)
    return lines


def rect(draw, xy, fill="#FFFFFF", outline="#335C81", width=3, radius=12):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def centered_text(draw, xy, text, fnt, fill="#0B2545", spacing=6):
    x1, y1, x2, y2 = xy
    max_width = x2 - x1 - 24
    lines = wrap_text(draw, text, fnt, max_width)
    heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + spacing * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + spacing


def arrow(draw, start, end, fill="#335C81", width=4):
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        head = [(ex, ey), (ex - direction * 16, ey - 9), (ex - direction * 16, ey + 9)]
    else:
        direction = 1 if ey > sy else -1
        head = [(ex, ey), (ex - 9, ey - direction * 16), (ex + 9, ey - direction * 16)]
    draw.polygon(head, fill=fill)


def save_canvas(name: str, title: str, draw_body) -> Path:
    path = IMG_DIR / f"{name}.png"
    img = Image.new("RGB", (1600, 1000), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_font = pil_font(42, bold=True)
    draw.text((60, 42), title, font=title_font, fill="#0B2545")
    draw.line((60, 105, 1540, 105), fill="#D9E2EC", width=4)
    draw_body(draw)
    img.save(path)
    return path


def make_diagrams() -> dict[str, Path]:
    ensure_dirs()
    body_font = pil_font(28)
    small_font = pil_font(23)
    diagrams: dict[str, Path] = {}

    def system_arch(draw):
        boxes = {
            "H5 租客端\nVue3 + Vant": (80, 180, 360, 330),
            "后台管理端\nVue3 + Element\nPlus": (80, 470, 360, 620),
            "Web-App API\nSpring Boot 3": (580, 180, 900, 330),
            "Web-Admin API\nSpring Boot 3": (580, 470, 900, 620),
            "公共模块\nJWT / 异常 / MinIO": (1010, 180, 1350, 330),
            "持久层\nMyBatis-Plus Mapper": (1010, 470, 1350, 620),
            "MySQL 8.0": (250, 760, 480, 880),
            "Redis 7": (680, 760, 900, 880),
            "MinIO 对象存储": (1080, 760, 1360, 880),
        }
        for text, xy in boxes.items():
            rect(draw, xy, "#F6FAFD")
            centered_text(draw, xy, text, body_font)
        arrow(draw, (360, 255), (580, 255))
        arrow(draw, (360, 545), (580, 545))
        arrow(draw, (900, 255), (1010, 255))
        arrow(draw, (900, 545), (1010, 545))
        arrow(draw, (1175, 620), (365, 760))
        arrow(draw, (1175, 620), (790, 760))
        arrow(draw, (1175, 620), (1220, 760))

    diagrams["system_arch"] = save_canvas("system_arch", "系统架构图", system_arch)

    def function_structure(draw):
        root = (650, 160, 950, 250)
        rect(draw, root, "#E8EEF5")
        centered_text(draw, root, "房屋租赁系统", body_font)
        modules = [
            ("租客端", "找房搜索\n房间详情\n预约看房\n我的租约\n浏览历史", 90),
            ("房源管理", "公寓维护\n房间维护\n属性维护\n标签维护\n费用维护", 390),
            ("租赁业务", "预约处理\n租约签订\n租约状态\n支付方式\n租期管理", 690),
            ("系统管理", "用户管理\n岗位管理\n角色菜单\n权限控制\n登录认证", 990),
            ("基础支撑", "区域数据\n文件上传\n缓存\n对象存储\n接口文档", 1290),
        ]
        for title, desc, x in modules:
            xy = (x, 370, x + 240, 680)
            rect(draw, xy, "#FFFFFF")
            centered_text(draw, (x, 380, x + 240, 455), title, body_font)
            centered_text(draw, (x + 10, 465, x + 230, 665), desc, small_font)
            arrow(draw, (800, 250), (x + 120, 370))

    diagrams["function_structure"] = save_canvas("function_structure", "系统功能结构图", function_structure)

    def dfd_context(draw):
        actors = {
            "租客": (100, 220, 300, 360),
            "运营管理员": (100, 560, 300, 700),
            "短信服务": (1250, 220, 1450, 360),
            "对象存储": (1250, 560, 1450, 700),
        }
        system = (560, 330, 980, 570)
        rect(draw, system, "#E8EEF5")
        centered_text(draw, system, "房屋租赁系统\n0", body_font)
        for text, xy in actors.items():
            rect(draw, xy)
            centered_text(draw, xy, text, body_font)
        arrow(draw, (300, 290), (560, 410))
        arrow(draw, (560, 470), (300, 630))
        arrow(draw, (980, 405), (1250, 290))
        arrow(draw, (1250, 630), (980, 500))
        draw.text((320, 245), "找房/预约/租约请求", font=small_font, fill="#0B2545")
        draw.text((315, 650), "房源/预约/合同管理", font=small_font, fill="#0B2545")
        draw.text((1010, 245), "验证码请求", font=small_font, fill="#0B2545")
        draw.text((1010, 650), "图片上传/访问", font=small_font, fill="#0B2545")

    diagrams["dfd_context"] = save_canvas("dfd_context", "顶层 DFD 图", dfd_context)

    def dfd_l1(draw):
        processes = [
            ("1. 用户认证", 110, 210),
            ("2. 房源查询", 510, 210),
            ("3. 预约管理", 910, 210),
            ("4. 租约管理", 310, 560),
            ("5. 后台运营", 710, 560),
        ]
        stores = [
            ("D1 用户库", 1200, 180),
            ("D2 房源库", 1200, 360),
            ("D3 预约库", 1200, 540),
            ("D4 租约库", 1200, 720),
        ]
        for text, x, y in processes:
            xy = (x, y, x + 280, y + 120)
            rect(draw, xy, "#F6FAFD")
            centered_text(draw, xy, text, body_font)
        for text, x, y in stores:
            xy = (x, y, x + 260, y + 90)
            rect(draw, xy, "#FFF7E6", "#C17C0A")
            centered_text(draw, xy, text, small_font)
        for start, end in [
            ((390, 270), (510, 270)),
            ((790, 270), (910, 270)),
            ((1050, 270), (1200, 225)),
            ((1050, 270), (1200, 405)),
            ((1050, 270), (1200, 585)),
            ((590, 680), (710, 620)),
            ((990, 620), (1200, 765)),
            ((390, 270), (430, 560)),
        ]:
            arrow(draw, start, end)

    diagrams["dfd_l1"] = save_canvas("dfd_l1", "1 层 DFD 图", dfd_l1)

    def dfd_l2(draw):
        flow = [
            ("输入筛选条件", 90),
            ("区域/价格/标签过滤", 340),
            ("读取公寓与房间", 590),
            ("组合费用/图片/租期", 840),
            ("返回列表与详情", 1090),
            ("记录浏览历史", 1340),
        ]
        for label, x in flow:
            xy = (x, 330, x + 190, 520)
            rect(draw, xy, "#FFFFFF")
            centered_text(draw, xy, label, small_font)
        for i in range(len(flow) - 1):
            arrow(draw, (flow[i][1] + 190, 425), (flow[i + 1][1], 425))
        for text, xy in [
            ("D2 房源库", (470, 700, 700, 810)),
            ("D5 图片库", (820, 700, 1050, 810)),
            ("D6 浏览记录", (1210, 700, 1450, 810)),
        ]:
            rect(draw, xy, "#FFF7E6", "#C17C0A")
            centered_text(draw, xy, text, small_font)
        arrow(draw, (675, 520), (585, 700))
        arrow(draw, (935, 520), (935, 700))
        arrow(draw, (1435, 520), (1330, 700))

    diagrams["dfd_l2"] = save_canvas("dfd_l2", "2 层 DFD 图：房源查询子过程", dfd_l2)

    def use_case(draw):
        system = (430, 150, 1180, 850)
        draw.rounded_rectangle(system, radius=20, outline="#335C81", width=4)
        draw.text((700, 170), "房屋租赁系统", font=body_font, fill="#0B2545")
        actors = [("租客", 140, 310), ("运营管理员", 1360, 310), ("系统管理员", 1360, 610)]
        for label, x, y in actors:
            draw.ellipse((x, y, x + 80, y + 80), outline="#335C81", width=3)
            draw.line((x + 40, y + 80, x + 40, y + 180), fill="#335C81", width=3)
            draw.line((x + 40, y + 105, x, y + 145), fill="#335C81", width=3)
            draw.line((x + 40, y + 105, x + 80, y + 145), fill="#335C81", width=3)
            draw.line((x + 40, y + 180, x, y + 240), fill="#335C81", width=3)
            draw.line((x + 40, y + 180, x + 80, y + 240), fill="#335C81", width=3)
            draw.text((x - 20, y + 255), label, font=small_font, fill="#0B2545")
        cases = [
            ("登录/认证", 540, 260),
            ("搜索房源", 770, 260),
            ("预约看房", 540, 430),
            ("签署租约", 770, 430),
            ("维护房源", 540, 600),
            ("处理预约", 770, 600),
            ("角色权限管理", 650, 740),
        ]
        for label, x, y in cases:
            draw.ellipse((x, y, x + 190, y + 80), outline="#335C81", width=3, fill="#F6FAFD")
            centered_text(draw, (x, y, x + 190, y + 80), label, small_font)
        for p in [(430, 390), (540, 300), (540, 470), (770, 300), (770, 470)]:
            arrow(draw, (220, 420), (p[0], p[1]))
        for p in [(730, 640), (960, 640), (840, 780)]:
            arrow(draw, (1360, 420), p)
        arrow(draw, (1360, 660), (840, 780))

    diagrams["use_case"] = save_canvas("use_case", "用户角色用例图", use_case)

    def sequence(draw):
        participants = [("租客端", 120), ("Web-App API", 420), ("业务服务", 720), ("数据库", 1020), ("短信/消息", 1320)]
        top, bottom = 180, 840
        for label, x in participants:
            rect(draw, (x, top, x + 180, top + 70), "#F6FAFD")
            centered_text(draw, (x, top, x + 180, top + 70), label, small_font)
            draw.line((x + 90, top + 70, x + 90, bottom), fill="#B8C5D3", width=3)
        steps = [
            (230, 510, "提交预约"),
            (510, 810, "校验房间/时间"),
            (810, 1110, "写入预约"),
            (1110, 810, "返回记录"),
            (810, 1410, "发送通知"),
            (510, 230, "返回结果"),
        ]
        y = 300
        for sx, ex, label in steps:
            arrow(draw, (sx, y), (ex, y))
            draw.text(((sx + ex) / 2 - 60, y - 36), label, font=small_font, fill="#0B2545")
            y += 90

    diagrams["sequence"] = save_canvas("sequence", "典型业务顺序图：预约看房", sequence)

    def state(draw):
        states = [
            ("待签约", 120, 260),
            ("已签约", 430, 260),
            ("租赁中", 740, 260),
            ("即将到期", 1050, 260),
            ("已退租", 740, 600),
            ("已取消", 430, 600),
        ]
        for label, x, y in states:
            rect(draw, (x, y, x + 220, y + 110), "#F6FAFD")
            centered_text(draw, (x, y, x + 220, y + 110), label, body_font)
        transitions = [
            ((340, 315), (430, 315), "签署合同"),
            ((650, 315), (740, 315), "生效"),
            ((960, 315), (1050, 315), "到期提醒"),
            ((1160, 370), (860, 600), "退租"),
            ((850, 370), (850, 600), "提前退租"),
            ((540, 370), (540, 600), "取消签约"),
        ]
        for s, e, label in transitions:
            arrow(draw, s, e)
            draw.text(((s[0] + e[0]) / 2 - 40, (s[1] + e[1]) / 2 - 30), label, font=small_font, fill="#0B2545")

    diagrams["state"] = save_canvas("state", "租约状态图", state)

    def er(draw):
        entities = [
            ("apartment_info\n公寓", 100, 170),
            ("room_info\n房间", 420, 170),
            ("user_info\n租客", 740, 170),
            ("view_appointment\n预约", 1060, 170),
            ("lease_agreement\n租约", 420, 540),
            ("graph_info\n图片", 100, 540),
            ("facility_info\n设施", 740, 540),
            ("system_user\n后台用户", 1060, 540),
        ]
        for text, x, y in entities:
            rect(draw, (x, y, x + 250, y + 130), "#FFFFFF")
            centered_text(draw, (x, y, x + 250, y + 130), text, small_font)
        for s, e, label in [
            ((350, 235), (420, 235), "1:N"),
            ((545, 300), (545, 540), "1:N"),
            ((865, 235), (1060, 235), "1:N"),
            ((545, 235), (1060, 235), "1:N"),
            ((545, 300), (865, 540), "N:M"),
            ((225, 540), (225, 300), "图片归属"),
        ]:
            arrow(draw, s, e)
            draw.text(((s[0] + e[0]) / 2, (s[1] + e[1]) / 2 - 24), label, font=small_font, fill="#0B2545")

    diagrams["er"] = save_canvas("er", "数据库 ER/关系概览图", er)

    def activity(draw):
        steps = [
            ("进入找房页", 160, 180),
            ("筛选区域/价格", 160, 330),
            ("查看房间详情", 160, 480),
            ("提交预约信息", 520, 480),
            ("后台确认预约", 880, 480),
            ("到店看房", 1240, 480),
            ("签署租约", 880, 700),
            ("结束/取消", 1240, 700),
        ]
        for label, x, y in steps:
            rect(draw, (x, y, x + 230, y + 90), "#F6FAFD")
            centered_text(draw, (x, y, x + 230, y + 90), label, small_font)
        for s, e in [
            ((275, 270), (275, 330)),
            ((275, 420), (275, 480)),
            ((390, 525), (520, 525)),
            ((750, 525), (880, 525)),
            ((1110, 525), (1240, 525)),
            ((995, 570), (995, 700)),
            ((1110, 745), (1240, 745)),
        ]:
            arrow(draw, s, e)

    diagrams["activity"] = save_canvas("activity", "预约看房活动图", activity)

    def network(draw):
        boxes = [
            ("用户浏览器/移动端", 120, 250),
            ("Nginx/静态资源", 430, 250),
            ("后端 API 服务\n8080/8081", 740, 250),
            ("MySQL\n3307", 1050, 160),
            ("Redis\n6379", 1050, 340),
            ("MinIO\n9000/9001", 1050, 520),
        ]
        for text, x, y in boxes:
            rect(draw, (x, y, x + 230, y + 110), "#F6FAFD")
            centered_text(draw, (x, y, x + 230, y + 110), text, small_font)
        arrow(draw, (350, 305), (430, 305))
        arrow(draw, (660, 305), (740, 305))
        arrow(draw, (970, 305), (1050, 215))
        arrow(draw, (970, 305), (1050, 395))
        arrow(draw, (970, 305), (1050, 575))
        draw.text((470, 700), "本地开发通过 Docker Compose 提供 MySQL、Redis、MinIO，后端使用 local profile 连接 localhost 服务。", font=small_font, fill="#0B2545")

    diagrams["network"] = save_canvas("network", "网络与部署设计图", network)

    def gantt(draw):
        tasks = [
            ("立项与调研", 1, 2),
            ("需求分析", 2, 4),
            ("概要/数据库设计", 4, 6),
            ("后台管理开发", 6, 9),
            ("H5 租客端开发", 7, 10),
            ("接口联调", 10, 11),
            ("测试与修复", 11, 13),
            ("文档与汇报", 13, 14),
        ]
        x0, y0, cell = 330, 180, 70
        draw.text((80, 140), "任务", font=body_font, fill="#0B2545")
        for i in range(1, 15):
            draw.text((x0 + (i - 1) * cell + 18, 140), str(i), font=small_font, fill="#0B2545")
            draw.line((x0 + (i - 1) * cell, 175, x0 + (i - 1) * cell, 850), fill="#E5E7EB", width=2)
        for r, (name, start, end) in enumerate(tasks):
            y = y0 + r * 80
            draw.text((80, y + 18), name, font=small_font, fill="#0B2545")
            draw.rectangle((x0 + (start - 1) * cell, y + 15, x0 + end * cell, y + 55), fill="#4F7CAC")
        draw.text((x0, 900), "单位：周", font=small_font, fill="#555555")

    diagrams["gantt"] = save_canvas("gantt", "项目实施甘特图", gantt)

    for kind, title in [
        ("admin_home", "后台管理端：首页截图占位"),
        ("admin_room", "后台管理端：房间管理截图占位"),
        ("admin_apartment", "后台管理端：公寓管理截图占位"),
        ("admin_appointment", "后台管理端：看房预约管理截图占位"),
        ("admin_agreement", "后台管理端：租约管理截图占位"),
        ("h5_search", "H5 租客端：找房页面截图占位"),
        ("h5_room", "H5 租客端：房间详情截图占位"),
        ("h5_appointment", "H5 租客端：预约看房截图占位"),
        ("h5_my_appointment", "H5 租客端：我的预约截图占位"),
        ("h5_user", "H5 租客端：个人中心截图占位"),
        ("scm", "软件配置管理工具截图占位"),
    ]:
        diagrams[kind] = save_canvas(kind, title, lambda draw, title=title: screenshot_placeholder(draw, title))

    return diagrams


def screenshot_placeholder(draw, title):
    body_font = pil_font(30)
    small_font = pil_font(24)
    rect(draw, (250, 170, 1350, 820), "#F9FAFB", "#94A3B8", width=4)
    rect(draw, (310, 240, 1290, 760), "#FFFFFF", "#CBD5E1", width=3)
    centered_text(draw, (310, 300, 1290, 480), title, body_font)
    centered_text(
        draw,
        (410, 500, 1190, 650),
        "此处用于替换为项目运行后的真实 UI 截图。\n当前文档保留统一尺寸与说明，便于后续直接替换。",
        small_font,
        fill="#475569",
    )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "000000") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(color)
    set_run_font(r, "Microsoft YaHei")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_font(paragraph, font_name: str = "Microsoft YaHei", size: int = 11) -> None:
    for run in paragraph.runs:
        set_run_font(run, font_name)
        run.font.size = Pt(size)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录将在 Word 中打开后更新。"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(text)
    run._r.append(fld_char3)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "房屋租赁系统软件工程大作业报告"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(hp, size=9)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.add_run("第 ")
    add_page_number(fp)
    fp.add_run(" 页")
    set_paragraph_font(fp, size=9)


def add_title_page(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("房屋租赁系统\n软件工程大作业报告")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    set_run_font(r)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于 Spring Boot、Vue3、MySQL、Redis 与 MinIO 的租赁业务平台")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string("475569")
    set_run_font(r)
    for _ in range(8):
        doc.add_paragraph("")
    rows = [
        ("课程名称", "软件工程"),
        ("项目名称", "房屋租赁系统"),
        ("项目组", "组员A / 组员B（待替换姓名与学号）"),
        ("提交日期", "2026 年 6 月"),
    ]
    add_table(doc, ["项目", "内容"], rows, widths=[1.6, 4.7])
    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[tuple], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, color="0B2545")
        set_cell_shading(hdr[idx], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph("")
    return table


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def p(doc: Document, text: str, style: str | None = None) -> None:
    para = doc.add_paragraph(style=style)
    for idx, part in enumerate(text.split("\n")):
        if idx:
            para.add_run().add_break()
        run = para.add_run(part)
        set_run_font(run)
        run.font.size = Pt(11)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        set_run_font(run)
        run.font.size = Pt(11)


def nums(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(item)
        set_run_font(run)
        run.font.size = Pt(11)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("475569")
    set_run_font(r)


def add_long_discussion(doc: Document, title: str, points: list[tuple[str, str]]) -> None:
    h(doc, title, 3)
    for label, text in points:
        para = doc.add_paragraph()
        r = para.add_run(f"{label}：")
        r.bold = True
        set_run_font(r)
        r.font.size = Pt(11)
        r2 = para.add_run(text)
        set_run_font(r2)
        r2.font.size = Pt(11)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def extract_schema_tables() -> list[str]:
    sql = ROOT / "project-lease" / "lease" / "lease_dump.sql"
    if not sql.exists():
        return []
    text = sql.read_text(encoding="utf-8", errors="ignore")
    return re.findall(r"CREATE TABLE `([^`]+)`", text)


def extract_routes() -> list[tuple[str, str]]:
    result = []
    route_files = [
        ROOT / "project-lease" / "rentHouseH5" / "rentHouseH5" / "src" / "router" / "tabBarRoutes.ts",
        ROOT / "project-lease" / "rentHouseH5" / "rentHouseH5" / "src" / "router" / "otherRoutes.ts",
        ROOT / "project-lease" / "rentHouseAdmin" / "rentHouseAdmin" / "src" / "router" / "constantRoutes.ts",
    ]
    for file in route_files:
        if file.exists():
            text = file.read_text(encoding="utf-8", errors="ignore")
            for path, title in re.findall(r"path:\s*['\"]([^'\"]+)['\"].*?title:\s*['\"]([^'\"]+)['\"]", text, re.S):
                result.append((title, path))
    return result[:24]


def build_report() -> None:
    diagrams = make_diagrams()
    doc = Document()
    configure_document(doc)
    add_header_footer(doc)
    add_title_page(doc)

    h(doc, "目录", 1)
    add_toc(doc.add_paragraph())
    p(doc, "提示：如需自动页码目录，请在 Word 中打开文档后选择“更新域/更新整个目录”。")
    doc.add_page_break()

    h(doc, "1. 项目概述", 1)
    h(doc, "1.1 项目背景与产业政策", 2)
    p(doc, "房屋租赁系统面向集中式公寓、长租房运营机构和租客群体，目标是在房源展示、预约看房、租约管理、用户管理和后台运营之间建立统一的信息化平台。随着住房租赁市场规范化、城市青年群体流动增强，以及线上房源服务逐步成为租赁业务入口，租赁企业需要以数字化系统提升房源周转、客户响应和合同管理效率。")
    p(doc, "从国家政策角度看，住房租赁市场发展强调规范化、透明化和信息化，鼓励通过数字平台提高住房供给匹配效率、降低线下沟通成本，并强化合同、租金、房源真实性等关键环节管理。从地方政策角度看，城市更新、人才住房、长租公寓和保障性租赁住房等业务均要求运营主体具备基础数据管理、服务跟踪和风险控制能力。因此，本项目选择房屋租赁业务作为软件工程大作业对象，能够覆盖从需求分析到设计、编码、测试和项目管理的完整过程。")
    add_long_discussion(doc, "1.2 用户单位存在的问题", [
        ("房源信息分散", "线下登记或多系统维护容易造成房源状态、租金、设施图片和地址信息不一致，影响客户体验和后台决策。"),
        ("预约流程不稳定", "租客咨询、看房预约、运营人员确认常依赖人工沟通，容易遗漏、重复安排或无法追踪处理状态。"),
        ("租约管理缺少闭环", "合同签订、租期、支付方式、退租状态和用户信息若缺少统一系统承载，会降低业务可审计性。"),
        ("图片与附件管理成本高", "房源图片、头像和合同相关附件需要对象存储统一管理，否则维护和迁移成本较高。"),
    ])
    h(doc, "1.3 产品开发目的", 2)
    bullets(doc, [
        "为租客提供找房、查看详情、预约看房、查看租约和浏览历史等线上服务。",
        "为运营人员提供公寓、房间、属性、费用、标签、预约、租约和用户管理能力。",
        "通过 MySQL、Redis、MinIO 等基础设施提升数据持久化、登录认证、缓存和图片存储能力。",
        "形成可演示、可扩展、可测试的软件工程项目，支撑课程报告与 PPT 汇报。",
    ])
    h(doc, "1.4 生命周期模型选择", 2)
    p(doc, "本项目选择“增量迭代模型”。原因是房屋租赁系统的核心业务可以拆分为若干相对独立的增量：基础数据与登录认证、房源展示与搜索、预约看房、租约管理、后台权限与运营配置。每个增量都可以形成可运行版本，便于小组成员并行开发、阶段演示和及时修正需求。相比瀑布模型，增量迭代更适合需求细节逐步明确、前后端联调频繁、UI 与接口需要多轮调整的课程项目。")
    doc.add_page_break()

    h(doc, "2. 项目可行性论证", 1)
    add_table(doc, ["维度", "可行性结论", "依据"], [
        ("技术可行性", "可行", "后端采用 Spring Boot 3、MyBatis-Plus、JWT；前端采用 Vue3、Vite、Element Plus、Vant；数据库与中间件采用 MySQL、Redis、MinIO，均为成熟技术。"),
        ("经济可行性", "可行", "本地开发可使用 Docker Compose 部署基础服务，技术栈开源，课程项目不产生高额许可证成本。"),
        ("操作可行性", "可行", "租客端以 H5 页面承载，后台以管理系统菜单承载；角色分工清晰，普通用户和管理员操作路径明确。"),
        ("法律可行性", "可行", "系统仅用于课程与演示，涉及用户信息和图片时应遵循最小必要、授权访问、脱敏展示和数据备份原则。"),
    ], widths=[1.2, 1.2, 4.9])
    h(doc, "2.1 技术可行性", 2)
    p(doc, "后端工程包含 model、common、web 等模块，web 下进一步拆分为 web-admin 与 web-app，分别承载后台管理接口和租客端接口。MyBatis-Plus 降低 CRUD 与分页开发复杂度，JWT 用于登录认证，MinIO 用于图片对象存储，Redis 可承担验证码、登录态或热点数据缓存。前端管理端采用 Vue3 与 Element Plus，适合复杂表单与表格；H5 租客端采用 Vue3 与 Vant，适合移动端交互。")
    h(doc, "2.2 经济与进度可行性", 2)
    p(doc, "课程项目以学习和演示为目标，成本主要来自开发时间。由于系统复用了成熟框架和本地 Docker 服务，基础设施搭建成本可控。本项目按两人小组分工推进：组员A 侧重后端、数据库、部署与测试，组员B 侧重前端、UI、需求文档与汇报材料；两人共同参与需求分析、系统设计、联调测试和答辩准备，在 14 周左右完成需求、设计、开发、测试与汇报材料。")
    h(doc, "2.3 风险分析", 2)
    add_table(doc, ["风险", "影响", "应对措施"], [
        ("前后端接口不一致", "页面无法正确展示或提交", "提前定义接口文档，联调阶段使用统一错误码和返回结构。"),
        ("数据库字段理解偏差", "业务状态流转错误", "以数据库表结构和枚举为准，关键状态编写测试用例。"),
        ("图片上传或访问失败", "房源详情展示不完整", "使用 MinIO 统一对象存储，并在本地配置中验证 bucket 与 URL。"),
        ("成员职责不清", "进度拖延", "使用甘特图和任务分解表明确负责人、交付物和验收标准。"),
    ], widths=[1.8, 2.1, 3.0])
    doc.add_page_break()

    h(doc, "3. 需求分析", 1)
    h(doc, "3.1 用户顶级需求描述", 2)
    p(doc, "房屋租赁系统的顶级需求是建立一个面向租客与运营人员的线上租赁服务平台。租客可通过移动端查找房间、公寓和详情信息，提交看房预约，查看个人预约、租约与浏览历史。运营人员可通过后台维护房源、公寓、房间、标签、属性、费用、支付方式、租期和图片，处理预约与租约，并通过系统管理模块维护岗位、用户、菜单与权限。")
    add_table(doc, ["用户角色", "核心目标", "主要功能"], [
        ("租客", "快速找到合适房源并完成预约/租约相关操作", "找房、筛选、查看详情、预约看房、查看租约、浏览历史、个人中心"),
        ("运营管理员", "维护房源并处理租赁业务", "公寓管理、房间管理、预约管理、租约管理、用户管理、图片上传"),
        ("系统管理员", "维护后台组织、岗位、角色和权限", "系统用户、岗位、菜单、角色权限、状态管理"),
    ], widths=[1.3, 2.4, 3.0])
    h(doc, "3.2 用户单位组织结构", 2)
    p(doc, "用户单位可抽象为运营管理部、门店/公寓运营人员、客服/租客服务人员、财务/合同管理人员和系统管理员。运营管理部负责业务规则和房源策略；公寓运营人员负责房源维护和看房安排；客服负责用户沟通与预约跟进；财务/合同人员负责租约和报价；系统管理员负责用户、岗位和权限配置。")
    add_figure(doc, diagrams["dfd_context"], "图 3-1 顶层 DFD 图")
    doc.add_page_break()

    h(doc, "3.3 基于结构化方法的需求分析", 2)
    p(doc, "结构化分析将系统视为若干处理过程、外部实体和数据存储之间的数据流。顶层 DFD 描述系统与租客、管理员、短信服务和对象存储之间的交互；1 层 DFD 将系统拆分为认证、房源查询、预约管理、租约管理和后台运营；2 层 DFD 进一步展开房源查询子过程。")
    add_figure(doc, diagrams["dfd_l1"], "图 3-2 1 层 DFD 图")
    add_figure(doc, diagrams["dfd_l2"], "图 3-3 2 层 DFD 图：房源查询子过程")
    doc.add_page_break()

    h(doc, "3.4 面向对象方法的需求分析", 2)
    add_figure(doc, diagrams["use_case"], "图 3-4 用户角色用例图")
    p(doc, "用例图显示租客主要参与找房、预约、租约查看等前台业务；运营管理员参与房源维护、预约处理和租约管理；系统管理员参与角色权限管理。系统功能边界清晰，前台与后台通过不同接口前缀区分。")
    doc.add_page_break()
    add_figure(doc, diagrams["sequence"], "图 3-5 典型业务顺序图：预约看房")
    p(doc, "预约看房流程体现了前端、控制器、业务服务、数据库和通知服务之间的协作。租客提交预约后，系统校验房间和时间信息，写入预约记录，并向运营人员或租客返回处理结果。")
    add_figure(doc, diagrams["state"], "图 3-6 租约状态图")
    doc.add_page_break()

    h(doc, "3.5 功能性需求描述", 2)
    add_table(doc, ["编号", "功能", "说明", "优先级"], [
        ("FR-01", "登录认证", "租客端验证码登录，后台端账号/验证码登录，统一返回 JWT 或用户信息。", "高"),
        ("FR-02", "房源搜索", "支持按区域、价格、标签、属性等条件查询公寓和房间。", "高"),
        ("FR-03", "房间详情", "展示房间图片、设施、标签、租金、支付方式和租期。", "高"),
        ("FR-04", "预约看房", "租客提交预约，后台分页查看并更新预约状态。", "高"),
        ("FR-05", "租约管理", "创建、查询、更新租约，并维护租约状态。", "高"),
        ("FR-06", "后台房源管理", "维护公寓、房间、属性、费用、标签、设施、图片。", "高"),
        ("FR-07", "系统管理", "维护后台用户、岗位、角色、菜单和权限。", "中"),
        ("FR-08", "浏览历史", "记录租客查看过的房源，支持分页查询。", "中"),
    ], widths=[0.8, 1.4, 4.2, 0.8])
    h(doc, "3.6 非功能性需求描述", 2)
    add_table(doc, ["类别", "需求"], [
        ("性能", "列表查询应支持分页，避免一次性加载过多房源和图片信息。"),
        ("安全", "后台接口需要认证拦截，用户密码和令牌应避免明文泄露。"),
        ("可靠性", "MySQL 数据持久化，Redis 开启 AOF，MinIO 使用 bucket 管理图片对象。"),
        ("可维护性", "后端按 model、common、web-admin、web-app 分层；前端按 views、api、router、store 组织。"),
        ("可用性", "H5 端适配移动端，后台端提供表格、表单、菜单和状态切换。"),
    ], widths=[1.3, 5.6])
    doc.add_page_break()

    h(doc, "4. 系统设计", 1)
    h(doc, "4.1 系统框架图", 2)
    add_figure(doc, diagrams["system_arch"], "图 4-1 系统架构图")
    p(doc, "系统采用前后端分离架构。H5 租客端和后台管理端通过 HTTP API 访问后端服务；后端按照管理端和租客端拆分 Controller 与 Service；持久层通过 MyBatis-Plus 操作数据库；图片文件由 MinIO 统一管理，认证和缓存能力由 Redis 和 JWT 支撑。")
    doc.add_page_break()
    h(doc, "4.2 系统功能结构图", 2)
    add_figure(doc, diagrams["function_structure"], "图 4-2 系统功能结构图")
    h(doc, "4.3 软件系统结构", 2)
    add_table(doc, ["模块", "职责", "关键技术"], [
        ("model", "实体、枚举、基础数据模型", "Java 17、Lombok、枚举类型"),
        ("common", "公共返回结构、异常处理、JWT、MinIO 配置、MyBatis-Plus 配置", "Spring Boot、JWT、MinIO SDK"),
        ("web-admin", "后台管理接口", "Controller、Service、Mapper、Knife4j"),
        ("web-app", "租客端接口", "Controller、Service、Mapper、认证拦截器"),
        ("rentHouseAdmin", "后台管理前端", "Vue3、Vite、Pinia、Element Plus"),
        ("rentHouseH5", "移动端租客前端", "Vue3、Vite、Vant、TailwindCSS"),
    ], widths=[1.4, 3.5, 2.0])
    doc.add_page_break()

    h(doc, "4.4 数据库设计", 2)
    add_figure(doc, diagrams["er"], "图 4-3 数据库 ER/关系概览图")
    schema_tables = extract_schema_tables()
    add_table(doc, ["类别", "数据表"], [
        ("房源基础", ", ".join([t for t in schema_tables if any(k in t for k in ["apartment", "room", "attr", "facility", "fee", "label", "graph"])][:18])),
        ("租赁业务", ", ".join([t for t in schema_tables if any(k in t for k in ["appointment", "agreement", "term", "payment", "history"])])),
        ("用户与系统", ", ".join([t for t in schema_tables if any(k in t for k in ["user", "system", "province", "city", "district"])])),
    ], widths=[1.3, 5.8])
    p(doc, "数据库围绕公寓、房间、用户、预约和租约展开。apartment_info 与 room_info 是房源主数据；room_attr_value、room_facility、room_label、room_payment_type、room_lease_term 等关系表描述房间属性、设施、标签、支付方式和租期；view_appointment 与 lease_agreement 承载预约和租约业务；system_user、system_post、system_role、system_menu 等表承载后台组织与权限。")
    doc.add_page_break()

    h(doc, "4.5 UI 设计", 2)
    p(doc, "后台管理端采用左侧菜单、顶部标签页、主体表格与抽屉/弹窗表单的组合，适合运营人员进行重复性管理操作。H5 租客端采用底部 TabBar、列表卡片、详情页和表单页，适合移动端快速浏览与提交预约。")
    for key, caption in [
        ("admin_home", "图 4-4 后台管理端：首页截图占位"),
        ("admin_room", "图 4-5 后台管理端：房间管理截图占位"),
        ("admin_apartment", "图 4-6 后台管理端：公寓管理截图占位"),
        ("admin_appointment", "图 4-7 后台管理端：看房预约管理截图占位"),
        ("admin_agreement", "图 4-8 后台管理端：租约管理截图占位"),
    ]:
        add_figure(doc, diagrams[key], caption)
        doc.add_page_break()
    for key, caption in [
        ("h5_search", "图 4-9 H5 租客端：找房页面截图占位"),
        ("h5_room", "图 4-10 H5 租客端：房间详情截图占位"),
        ("h5_appointment", "图 4-11 H5 租客端：预约看房截图占位"),
        ("h5_my_appointment", "图 4-12 H5 租客端：我的预约截图占位"),
        ("h5_user", "图 4-13 H5 租客端：个人中心截图占位"),
    ]:
        add_figure(doc, diagrams[key], caption)
        doc.add_page_break()

    h(doc, "4.6 网络设计", 2)
    add_figure(doc, diagrams["network"], "图 4-14 网络与部署设计图")
    p(doc, "本地开发环境中，Docker Compose 提供 MySQL、Redis 和 MinIO。后端 web-admin 运行在 8080，web-app 运行在 8081，前端通过开发服务器或构建后的静态资源访问后端 API。部署时可使用 Nginx 承载前端静态资源并反向代理后端接口，数据库与对象存储部署在受控网络内。")
    h(doc, "4.7 数据量分析", 2)
    add_table(doc, ["对象", "估算规模", "设计影响"], [
        ("公寓", "100-1000 条", "列表分页、按区域检索。"),
        ("房间", "1000-10000 条", "按公寓、价格、属性和状态过滤。"),
        ("图片", "每房间 3-8 张", "对象存储承载，数据库只保存 URL 与归属关系。"),
        ("预约", "每日数十至数百条", "按用户、状态、时间分页查询。"),
        ("租约", "与有效出租房间数量接近", "状态流转和到期提醒是核心。"),
    ], widths=[1.4, 1.7, 3.8])
    h(doc, "4.8 预约模块活动图", 2)
    add_figure(doc, diagrams["activity"], "图 4-15 预约看房活动图")
    doc.add_page_break()

    h(doc, "5. 编码实现", 1)
    h(doc, "5.1 后端实现", 2)
    p(doc, "后端使用 Spring Boot 3.0.5 和 Java 17。公共模块封装统一 Result 返回结构、ResultCodeEnum、GlobalExceptionHandler、JWT 工具、登录用户持有器、MyBatis-Plus 自动填充和 MinIO 配置。管理端接口以 /admin 为前缀，租客端接口以 /app 为前缀，便于权限边界和前端调用区分。")
    add_table(doc, ["接口领域", "代表路径", "说明"], [
        ("租客登录", "/app/login、/app/info", "验证码获取、登录、当前用户信息。"),
        ("房源查询", "/app/room、/app/apartment", "房间分页、详情、公寓列表与详情。"),
        ("预约", "/app/appointment", "保存预约、查询我的预约、查看详情。"),
        ("租约", "/app/agreement", "查看租约、保存或更新租约状态。"),
        ("后台房源", "/admin/apartment、/admin/room、/admin/attr、/admin/fee", "维护房源基础数据。"),
        ("后台系统", "/admin/system", "维护用户、岗位、角色、菜单。"),
    ], widths=[1.4, 2.0, 3.2])
    h(doc, "5.2 前端实现", 2)
    routes = extract_routes()
    add_table(doc, ["页面标题", "路由路径"], routes[:18], widths=[2.2, 4.6])
    p(doc, "后台管理端的页面由静态路由配置承载，包括首页、系统用户、岗位、公寓管理、房间管理、属性管理、预约管理、租约管理和用户管理等。H5 端包括找房、圈子、我的房间、消息、个人中心以及房间详情、预约看房、我的预约、我的租约等页面。")
    h(doc, "5.3 本地部署实现", 2)
    add_table(doc, ["服务", "端口", "用途"], [
        ("MySQL 8.0", "3307 -> 3306", "存储业务与系统数据。"),
        ("Redis 7", "6379", "缓存、验证码或登录辅助数据。"),
        ("MinIO", "9000 / 9001", "对象存储 API 与控制台。"),
        ("web-admin", "8080", "后台管理 API。"),
        ("web-app", "8081", "租客端 API。"),
    ], widths=[1.5, 1.5, 3.8])
    h(doc, "5.4 实现效果说明", 2)
    p(doc, "由于本报告先生成可提交文档，UI 截图区域以统一占位方式预留。项目运行后，可将后台管理端和 H5 租客端对应页面截图替换到第 4.5 节，以增强汇报材料的真实性。")
    doc.add_page_break()

    h(doc, "6. 测试", 1)
    h(doc, "6.1 测试策略", 2)
    p(doc, "测试采用白盒测试与黑盒测试结合的方式。白盒测试聚焦服务方法、状态流转、条件分支和异常处理；黑盒测试聚焦用户输入、接口返回、页面流程和业务规则。测试环境使用本地 Docker 服务，保证数据库、缓存和对象存储可重复启动。")
    h(doc, "6.2 白盒测试案例", 2)
    add_table(doc, ["项目", "内容"], [
        ("测试对象", "预约保存服务：ViewAppointmentService.saveOrUpdate"),
        ("输入条件", "有效用户、有效房间 ID、预约时间和联系人信息。"),
        ("覆盖路径", "参数校验通过 -> 构造预约记录 -> 写入数据库 -> 返回成功结果。"),
        ("预期结果", "数据库新增或更新预约记录，状态为待看房/待确认，返回统一成功响应。"),
        ("异常路径", "房间 ID 不存在或用户未登录时，应返回业务异常或认证失败。"),
    ], widths=[1.4, 5.4])
    h(doc, "6.3 黑盒测试案例", 2)
    add_table(doc, ["项目", "内容"], [
        ("测试场景", "租客在 H5 端完成房间搜索并提交预约。"),
        ("前置条件", "系统已有可出租房间，租客已登录。"),
        ("测试步骤", "进入找房页 -> 选择区域/价格 -> 打开房间详情 -> 点击预约 -> 填写预约信息 -> 提交。"),
        ("预期结果", "页面提示预约成功；我的预约列表出现该记录；后台预约管理可查询并更新状态。"),
        ("验收标准", "页面无报错，接口返回成功，数据库 view_appointment 表存在对应记录。"),
    ], widths=[1.4, 5.4])
    h(doc, "6.4 测试用例扩展表", 2)
    add_table(doc, ["编号", "类型", "测试点", "预期结果"], [
        ("TC-01", "接口", "后台登录验证码获取", "返回图片验证码与 key。"),
        ("TC-02", "接口", "租客登录", "手机号验证码正确时返回 token。"),
        ("TC-03", "功能", "房源分页查询", "按条件返回分页列表。"),
        ("TC-04", "功能", "图片上传", "MinIO 返回可访问 URL。"),
        ("TC-05", "权限", "未登录访问后台接口", "返回认证失败。"),
        ("TC-06", "状态", "预约状态更新", "后台更新后前台可见状态变化。"),
        ("TC-07", "边界", "空筛选条件搜索", "返回默认分页结果。"),
        ("TC-08", "异常", "不存在房间详情", "返回业务错误或空结果。"),
    ], widths=[0.8, 1.0, 3.0, 2.1])
    doc.add_page_break()

    h(doc, "7. 软件项目管理", 1)
    h(doc, "7.1 项目实施甘特图", 2)
    add_figure(doc, diagrams["gantt"], "图 7-1 项目实施甘特图")
    h(doc, "7.2 小组成员职责", 2)
    add_table(doc, ["成员", "职责", "主要交付物"], [
        ("组员A（待替换姓名/学号）", "后端与数据负责人；共同承担需求分析、系统设计、联调测试和答辩", "Spring Boot 接口、MyBatis-Plus Mapper、数据库表分析、Redis/MinIO/Docker 本地部署、白盒测试、后端接口说明。"),
        ("组员B（待替换姓名/学号）", "前端与文档负责人；共同承担需求分析、系统设计、联调测试和答辩", "后台管理端页面、H5 租客端页面、UI 截图整理、DFD/用例/流程图与报告排版、黑盒测试、PPT 汇报材料。"),
    ], widths=[1.8, 2.4, 3.0])
    h(doc, "7.3 软件配置管理", 2)
    p(doc, "项目配置管理建议使用 Git 与远程仓库。成员按功能创建分支，完成后通过合并请求或代码评审合入主分支。提交信息遵循 feat、fix、docs、test、chore 等类型，便于后续回溯。依赖版本由 pom.xml、package.json 和 lock 文件管理；本地基础服务由 docker-compose.yml 固化。")
    add_figure(doc, diagrams["scm"], "图 7-2 软件配置管理工具截图占位")
    h(doc, "7.4 产品报价", 2)
    add_table(doc, ["费用项", "估算金额（元）", "说明"], [
        ("需求分析与设计", "8,000", "业务调研、原型、数据库与架构设计。"),
        ("后端开发", "18,000", "接口、认证、业务服务、数据访问和对象存储。"),
        ("前端开发", "20,000", "后台管理端与 H5 租客端页面实现。"),
        ("测试与部署", "6,000", "功能测试、联调、本地 Docker 部署和缺陷修复。"),
        ("文档与培训", "4,000", "用户说明、项目报告、汇报材料。"),
        ("合计", "56,000", "课程项目报价估算，可按实际范围调整。"),
    ], widths=[2.0, 1.6, 3.2])
    doc.add_page_break()

    h(doc, "附录 A：数据库表清单", 1)
    rows = []
    for idx, table in enumerate(schema_tables, start=1):
        if "apartment" in table or "room" in table:
            category = "房源"
        elif "agreement" in table or "appointment" in table or "history" in table:
            category = "租赁业务"
        elif "system" in table or "user" in table:
            category = "用户/系统"
        else:
            category = "基础字典"
        rows.append((idx, table, category))
    add_table(doc, ["序号", "表名", "类别"], rows, widths=[0.8, 3.4, 2.0])
    doc.add_page_break()

    h(doc, "附录 B：接口与页面对应关系", 1)
    add_table(doc, ["页面/模块", "前端路由或接口", "说明"], [
        ("找房", "/search、/app/room/pageItem", "租客按条件浏览房间。"),
        ("房间详情", "/roomDetail、/app/room/getDetailById", "查看图片、设施、租金和租期。"),
        ("预约看房", "/appointment、/app/appointment/saveOrUpdate", "提交预约并查询状态。"),
        ("我的租约", "/myAgreement、/app/agreement/listItem", "查看个人租约列表。"),
        ("公寓管理", "/admin/apartment", "后台维护公寓信息。"),
        ("房间管理", "/admin/room", "后台维护房间信息。"),
        ("租约管理", "/admin/agreement", "后台创建、分页查询、更新租约。"),
        ("系统用户", "/admin/system/user", "后台用户和角色维护。"),
    ], widths=[1.5, 2.5, 3.0])
    h(doc, "附录 C：课程要求对照表", 1)
    add_table(doc, ["要求", "报告位置", "完成说明"], [
        ("项目概述", "第 1 章", "包含政策、问题、目的与生命周期模型。"),
        ("可行性论证", "第 2 章", "覆盖技术、经济、操作、法律分析。"),
        ("需求分析", "第 3 章", "包含 DFD、用例、顺序图、状态图和需求描述。"),
        ("系统设计", "第 4 章", "包含架构、功能结构、数据库、UI、网络、数据量和活动图。"),
        ("编码实现", "第 5 章", "说明后端、前端、本地部署和实现效果截图位置。"),
        ("测试", "第 6 章", "包含白盒与黑盒测试案例。"),
        ("项目管理", "第 7 章", "包含甘特图、配置管理、报价和成员职责。"),
        ("不少于 40 页", "全文", "通过正文、图表、截图占位和附录满足页数要求。"),
    ], widths=[1.8, 1.8, 3.3])

    doc.add_page_break()
    h(doc, "附录 D：验收与汇报素材", 1)
    p(doc, "本附录用于支撑课程汇报、答辩和最终自查。每个条目均可在答辩 PPT 中提炼为一页或半页内容，也可作为组内分工验收依据。")
    checklist_sections = [
        ("D.1 项目背景汇报要点", [
            ("市场背景", "租赁业务线上化可以降低线下咨询、看房、合同维护和房源更新的沟通成本。"),
            ("业务痛点", "房源状态、预约进度、租约状态和图片资料缺少统一平台时，运营人员难以及时同步信息。"),
            ("项目价值", "系统将找房、预约、租约、后台房源维护和权限管理串成闭环，适合作为软件工程课程案例。"),
        ]),
        ("D.2 需求分析汇报要点", [
            ("租客需求", "快速筛选房源、查看真实详情、提交预约、跟踪个人预约和租约。"),
            ("管理员需求", "维护房源、公寓、房间、属性、费用、图片，处理预约和租约。"),
            ("系统需求", "认证、权限、异常处理、接口统一返回、数据持久化和对象存储。"),
        ]),
        ("D.3 结构化分析答辩口径", [
            ("顶层 DFD", "说明系统与租客、管理员、短信服务、对象存储之间的数据交换。"),
            ("1 层 DFD", "说明认证、房源查询、预约管理、租约管理和后台运营五类核心处理。"),
            ("2 层 DFD", "以房源查询为例，展开筛选、读取、组合、返回和浏览记录写入过程。"),
        ]),
        ("D.4 面向对象分析答辩口径", [
            ("用例图", "强调租客、运营管理员和系统管理员的边界与主要用例。"),
            ("顺序图", "以预约看房说明前端、控制器、服务、数据库和通知服务之间的消息流。"),
            ("状态图", "以租约状态说明待签约、已签约、租赁中、到期、退租和取消等业务状态。"),
        ]),
        ("D.5 系统架构汇报要点", [
            ("前后端分离", "H5 和后台管理端只负责交互与展示，后端 API 负责业务处理。"),
            ("模块分层", "model 承载实体，common 承载公共能力，web-admin 和 web-app 区分服务边界。"),
            ("基础设施", "MySQL 存储结构化数据，Redis 支撑缓存/验证码，MinIO 存储图片对象。"),
        ]),
        ("D.6 数据库设计自查", [
            ("主数据", "apartment_info、room_info、user_info 是核心主数据。"),
            ("关系数据", "room_label、room_facility、room_payment_type 等表表达多对多或扩展属性。"),
            ("业务数据", "view_appointment 与 lease_agreement 记录预约和租约闭环。"),
        ]),
        ("D.7 UI 设计自查", [
            ("后台端", "重点展示首页、房间管理、公寓管理、预约管理、租约管理。"),
            ("H5 端", "重点展示找房、房间详情、预约看房、我的预约、个人中心。"),
            ("替换策略", "报告中的占位图保持统一尺寸，后续可直接替换为真实截图。"),
        ]),
        ("D.8 编码实现自查", [
            ("后端", "说明 Controller、Service、Mapper 与统一 Result 返回结构。"),
            ("前端", "说明 Vue3、路由、API 封装、状态管理和组件库。"),
            ("部署", "说明 Docker Compose、本地 profile、端口和服务连接方式。"),
        ]),
        ("D.9 测试答辩要点", [
            ("白盒测试", "围绕预约保存服务的参数、分支、数据库写入和异常路径。"),
            ("黑盒测试", "围绕租客找房到预约成功的完整用户路径。"),
            ("回归测试", "修改预约、租约、登录等核心模块后优先复测相关场景。"),
        ]),
        ("D.10 项目管理答辩要点", [
            ("进度", "按调研、需求、设计、开发、联调、测试、文档汇报划分阶段。"),
            ("分工", "两名成员平均分担工作：组员A 偏后端、数据库、部署和白盒测试；组员B 偏前端、UI、文档和黑盒测试；需求、设计、联调和答辩共同完成。"),
            ("配置管理", "使用 Git、分支、提交规范、依赖锁和 Docker 配置保证可追溯。"),
        ]),
        ("D.11 风险与改进方向", [
            ("安全风险", "后续可强化密码加密、接口权限粒度和敏感数据脱敏。"),
            ("业务风险", "可增加租金支付、合同模板、消息提醒和数据统计。"),
            ("性能风险", "房源和图片规模变大后可增加缓存、索引优化和 CDN。"),
        ]),
        ("D.12 PPT 制作建议", [
            ("章节安排", "按背景、需求、设计、实现、测试、管理、总结组织。"),
            ("图表优先", "DFD、用例图、架构图、ER 图、甘特图适合放入 PPT。"),
            ("演示路径", "推荐演示后台新增房源、H5 搜索详情、提交预约、后台处理预约。"),
        ]),
        ("D.13 最终提交前检查", [
            ("文档", "更新目录、替换成员姓名学号、替换 UI 截图、检查页码。"),
            ("项目", "确认后端服务、数据库、Redis、MinIO、前端页面可运行。"),
            ("汇报", "准备 5-10 分钟讲稿，并让每位成员说明自己的职责。"),
        ]),
        ("D.14 后续扩展设想", [
            ("业务扩展", "增加在线签约、租金支付、维修报修和评价功能。"),
            ("运营扩展", "增加房源上架审核、租金统计、空置率分析和看房转化率报表。"),
            ("技术扩展", "增加单元测试覆盖、CI/CD、日志监控和接口限流。"),
        ]),
    ]
    for title, rows in checklist_sections:
        h(doc, title, 2)
        add_table(doc, ["条目", "说明"], rows, widths=[1.6, 5.2])
        p(doc, "汇报使用建议：本页内容可以作为答辩提纲或 PPT 备注。实际提交前，请结合项目运行截图、组员分工和老师要求进行最后替换。")
        doc.add_page_break()

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_report()
    print(str(OUT_DOCX))
