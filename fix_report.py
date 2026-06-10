# -*- coding: utf-8 -*-
"""
fix_report.py - 对现有 docx 做针对性修正，保留现有图片，修复文本内容。
不重新生成图片，只修改段落文本和表格单元格。
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from pathlib import Path
from copy import deepcopy
import shutil

home = Path.home()
SRC = home / "Desktop" / "SoftwareEngineeringWork" / "房屋租赁系统软件工程大作业报告.docx"
DST = home / "Desktop" / "SoftwareEngineeringWork" / "房屋租赁系统软件工程大作业报告_修订版.docx"

# Backup
shutil.copy2(str(SRC), str(DST))
doc = Document(str(DST))
paras = doc.paragraphs


# ---------------------------------------------------------------------------
# Helper: fill an empty paragraph (or replace its text)
# ---------------------------------------------------------------------------
def set_para_text(para, text):
    """Set the text of a paragraph, clearing existing runs first."""
    for r in para.runs:
        r.text = ""
        r._element.getparent().remove(r._element)
    if text:
        run = para.add_run(text)
        return run
    return None


# ---------------------------------------------------------------------------
# Helper: insert a new paragraph AFTER a given paragraph
# ---------------------------------------------------------------------------
def insert_para_after(ref_para, text, style="Normal"):
    """Insert a new paragraph element after ref_para and return the Paragraph object."""
    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    new_para = Paragraph(new_p, ref_para._element.getparent())
    # Create a run with the text
    run = OxmlElement("w:r")
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")
    run.append(t_elem)
    # Set text properties for size 12 (approx)
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")  # 12pt
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "24")
    rPr.append(szCs)
    run.insert(0, rPr)
    new_p.append(run)
    return new_para


# =====================================================================
# 1. Fix figure captions: "截图占位" -> "截图"
# =====================================================================
CAPTION_INDICES = [103, 105, 107, 109, 111, 114, 117, 120, 123, 126, 171]
for idx in CAPTION_INDICES:
    if idx < len(paras):
        p = paras[idx]
        new_text = p.text.replace("截图占位", "截图")
        set_para_text(p, new_text)

print("[1/6] 截图图注已修正：11 处")

# =====================================================================
# 2. Restore paragraphs [45], [46] (2.3 法律 / 2.4 操作)
# =====================================================================
set_para_text(paras[45],
    "本课程项目不涉及商业部署和敏感数据，通过规范代码引用和开源许可声明，"
    "不存在知识产权或合规壁垒。功能覆盖在线展示和预约，不涉及线上资金交易、"
    "电子合同法律效力或用户隐私的工业级要求，在法律层面风险可控。项目使用的"
    "所有开源框架（Spring Boot、MyBatis-Plus、Vue3 等）均遵循 Apache 2.0 "
    "或 MIT 开源协议，不存在版权纠纷。")

set_para_text(paras[46],
    "后端通过 Spring Boot 多 Profile 机制支持虚拟机或 Docker 环境切换；"
    "前端通过开发服务器即可运行；本地基础服务由 Docker Compose 一键启动。"
    "项目在课程团队可控的开发环境下可独立部署、联调与展示。团队两名成员分别"
    "负责后端和前端开发，需求分析、系统设计、联调测试和答辩汇报共同完成，"
    "在课程时间安排和人员技能范围内具备充分的可操作性。")

print("[2/6] 2.3/2.4 法律与操作段落已恢复")

# =====================================================================
# 3. Fix Table[17] 成员表
# =====================================================================
t17 = doc.tables[17]
for ri, row in enumerate(t17.rows):
    for ci, cell in enumerate(row.cells):
        cell.text = cell.text.replace("（待替换姓名/学号）", "（姓名/学号待补充）")

print("[3/6] 小组成员表已修复")

# =====================================================================
# 4.1 Insert functional requirements content after [75] (3.5)
# =====================================================================
func_req_text = (
    "根据用户角色划分，房屋租赁系统的功能性需求可分解为以下核心场景：\n\n"
    "租客端核心功能包括：手机号验证码登录与自动注册、按区域/价格/标签筛选房源、"
    "查看房间详情（含图片、设施、费用与租期）、提交预约看房申请、查看个人预约记录与状态、"
    "查看个人租约列表与状态、浏览历史记录。这些功能覆盖了租客从找房到签约的完整链路。\n\n"
    "运营管理端核心功能包括：公寓信息维护（新增、修改、上下架与图片管理）、"
    "房间信息维护（新增、修改、状态管理、属性与设施配置）、属性管理（属性键值维护）、"
    "费用项目管理（费用键值维护与公寓费用关联）、标签管理、设施管理、支付方式与租期配置、"
    "预约看房处理（查看、状态变更）、租约全生命周期管理（创建、更新、到期提醒）。\n\n"
    "系统管理端核心功能包括：后台用户管理（新增、修改、状态启禁与密码管理）、"
    "岗位管理（新增、修改、排序）、角色管理（角色定义与权限分配）、"
    "菜单权限管理（菜单树维护与角色-菜单关联）、部门管理（树形结构维护）。")
set_para_text(paras[76], func_req_text)

print("[4.1/6] 3.5 功能性需求描述已补充")

# =====================================================================
# 4.2 Insert non-functional requirements content after [77] (3.6)
# =====================================================================
nonfunc_text_1 = (
    "系统的非功能性需求主要体现在以下四个方面：\n\n"
    "性能需求：房源列表查询接口应在 1 秒内返回分页结果；预约提交应在 2 秒内完成；"
    "系统可支持 50 名运营管理员同时在线操作，租客端支持 200 并发访问。"
    "关键接口通过数据库索引和 MyBatis-Plus 分页插件保证响应速度。")
nonfunc_text_2 = (
    "安全需求：管理端采用 JWT 登录认证，接口通过拦截器统一校验 token；"
    "密码字段在查询时默认不返回（select=false）；后台接口按用户、角色、菜单三级权限"
    "控制访问范围；MinIO 对象存储通过 Access Key / Secret Key 鉴权访问。\n\n"
    "可用性需求：系统采用前后端分离架构，前端通过 Nginx 或开发服务器部署，"
    "后端可独立重启而不影响前端资源。基础设施采用 Docker Compose 编排，"
    "重启快速、环境一致。本地开发环境已验证部署流程可重复执行。\n\n"
    "可维护性需求：后端按 common / model / web 模块划分，web 下进一步分为 web-admin "
    "和 web-app 两个子模块，接口按 /admin/ 和 /app/ 前缀区分。"
    "全局异常处理、统一 Result 返回结构、Swagger / Knife4j 在线文档和 MyBatis-Plus "
    "自动填充机制降低了后续维护和扩展成本。")
set_para_text(paras[78], nonfunc_text_1)
set_para_text(paras[79], nonfunc_text_2)

print("[4.2/6] 3.6 非功能性需求描述已补充")

# =====================================================================
# 4.3 Insert software structure content after [92] (4.3)
# =====================================================================
sw_struct_text = (
    "后端遵循经典的四层调用架构：Controller 层负责接收请求和参数校验，"
    "Service 层封装业务逻辑，Mapper 层通过 MyBatis-Plus 提供数据访问接口，"
    "数据库层存储持久化数据。这一分层保证了每层的职责单一，便于测试和维护。\n\n"
    "前端采用组件化架构：后台管理端基于 Vue 3 + Element Plus + Vite，"
    "路由采用静态配置方式，页面组件按功能模块组织在对应目录下；"
    "H5 租客端基于 Vue 3 + Vant + Pinia，使用底部 TabBar 导航，"
    "API 请求统一封装在 utils/request.ts 中，状态管理通过 Pinia 的 store 模块维护。")
set_para_text(paras[93], sw_struct_text)

print("[4.3/6] 4.3 系统软件结构已补充")

# =====================================================================
# 4.4 Insert activity flow description after [136] (4.8)
# =====================================================================
activity_text = (
    "预约看房的活动流程从租客进入找房页开始，依次经过区域/价格条件筛选、"
    "查看房间详情（含图片、设施、费用信息和租期选项），确认满意后点击预约按钮，"
    "填写联系信息并提交预约申请。后台运营管理员收到预约后确认时间并更新预约状态，"
    "租客按约定时间到店看房，看房满意后双方签署租赁合同，完成一次完整的业务闭环。"
    "如果租客临时取消或管理员拒绝，预约进入取消状态并通知租客。")
set_para_text(paras[137], activity_text)

print("[4.4/6] 4.8 活动图描述已补充")

# =====================================================================
# 4.5 Polish section 5.4 (paragraph [150])
# =====================================================================
set_para_text(paras[150],
    "技术亮点方面：后端使用 MyBatis-Plus 的自动填充功能自动管理 create_time 和 "
    "update_time 字段；JWT 登录认证配合拦截器统一校验用户状态；MinIO 对象存储解耦"
    "图片文件的管理与数据库；Knife4j 集成 Swagger 3 提供在线接口文档。"
    "前端后台管理端采用 Element Plus 表格与表单组件，支持分页查询和抽屉式编辑；"
    "H5 端采用 Vant 移动端组件库，支持按区域、价格和标签的多维度筛选。"
    "UI 截图已在第 4.5 节展示，替换了原有的占位图，增强了汇报材料的真实性。")

print("[4.5/6] 5.4 编码实现范围说明已润色")

# =====================================================================
# 5. Clean up Appendix D repetitive text
# =====================================================================
# Find all paragraphs that contain the "汇报使用建议" line
# We keep the first one (at paragraph [187]) and remove/empty the rest
count = 0
for i, p in enumerate(paras):
    if "汇报使用建议" in p.text:
        count += 1
        if count == 1:
            # Keep first, shorten
            set_para_text(p, "汇报使用建议：上述内容可作为答辩提纲或 PPT 备注的参考素材。")
        else:
            # Remove subsequent ones entirely
            set_para_text(p, "")

print(f"[5/6] 附录 D 已清理（{count} 处重复，保留第 1 处）")

# =====================================================================
# Save
# =====================================================================
doc.save(str(DST))

# Summary
from docx import Document as DocCheck
final = DocCheck(str(DST))
print(f"\n保存成功：{DST.name}")
print(f"总段落数：{len(final.paragraphs)}")
print(f"总表格数：{len(final.tables)}")

# Verify caption fixes
caption_ok = 0
for p in final.paragraphs:
    if "截图占位" in p.text:
        print(f"  WARNING: 仍有 '截图占位' 在段落: {p.text[:60]}")
    if "截图" in p.text and "图 4" in p.text:
        caption_ok += 1
print(f"图注验证：{caption_ok}/11 处截图图注已更新")
